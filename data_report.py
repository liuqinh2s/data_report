# coding=utf-8

import docx
from win32com import client
import yaml

from information_extraction import qichacha
from word_manipulation import word

f = open("data_report_config.yml", encoding="utf-8")
config = yaml.load(f, Loader=yaml.FullLoader)
f.close()


def replace_doc(doc_name):
    try:
        word = client.Dispatch("Word.Application")
        doc = word.Documents.Open(doc_name)
        # string--搜索文本
        # True--区分大小写,
        # True--完全匹配的单词，并非单词中的部分（全字匹配）,
        # True--使用通配符,
        # True--同音,
        # True--查找单词的各种形式,
        # True--向文档尾部搜索,
        # 1,
        # True--带格式的文本,
        # new_string--替换文本,
        # 2--替换个数（全部替换）
        for i in range(len(config['identity_info_replace_list'])):
            word.Selection.Find.Execute(config['identity_info_replace_list'][i], False, False, False, False, False,
                                        True, 1, True,
                                        BUSINESS_INFO[i], 2)
        doc.Close()
        word.Quit()
    except Exception as e:
        print(e)


def insert_table(docx_file, index, INFO):
    table = docx_file.tables[index]
    if len(INFO)==0:
        for i in range(len(table.rows)):
            row = table.add_row()
            for j in range(len(row.cells)):
                row.cells[j].text = '-'
    else:
        for i in range(len(INFO)):
            row = table.add_row()
            for j in range(len(row.cells)):
                row.cells[j].text = INFO[i][j]


def replace_docx(input_file, output_file, BUSINESS_INFO, INFO_LIST):
    try:
        docx_file = docx.Document(input_file)
        # 替换身份信息表格内容
        for row in docx_file.tables[0].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for i in range(len(config['identity_info_replace_list'])):
                        if config['identity_info_replace_list'][i] == paragraph.text:
                            paragraph.text = BUSINESS_INFO[i]
        for i in range(len(INFO_LIST)):
            insert_table(docx_file, i + 1, INFO_LIST[i])
        # 必须要保存才能修改生效
        docx_file.save(output_file[:-5]+"数据报告.docx")
    except Exception as e:
        print(e)


if __name__ == "__main__":
    doc_name = r"C:\Users\liuqinh2s\PycharmProjects\Python3\files\基础版企业信用报告—乐视控股（北京）有限公司—天眼查（W11032183931553137728089）.doc"
    docx_name = r"C:\Users\liuqinh2s\PycharmProjects\Python3\files\基础版企业信用报告—乐视控股（北京）有限公司—天眼查（W11032183931553137728089）.docx"
    data_report_docx = r"C:\Users\liuqinh2s\PycharmProjects\Python3\files\demo_yuanban200.docx"
    word.doc2docx(doc_name, docx_name)
    BUSINESS_INFO, INFO_LIST = qichacha.read_docx(docx_name)
    replace_docx(data_report_docx, docx_name, BUSINESS_INFO, INFO_LIST)
    print()
